import os
import glob
from docx import Document

#Om changed on Aug 25

# Define the source directory and destination file path for the first block
source_directory_1 = r'C:\Users\OMCENTER\projects\mytools\droidlogin\droid\droid'
destination_file = r'C:\Users\OMCENTER\projects\drOidwOrks_V_1_1\driodworks.doc'
## just new 

# Create a new Document or open an existing one
doc = Document()

# Find all .py files in the first source directory
py_files_1 = glob.glob(os.path.join(source_directory_1, '*.py'))

# Loop through each .py file in the first source directory
for py_file in py_files_1:
	# Add the file name before the content
	file_name = os.path.basename(py_file)
	doc.add_paragraph(f"<<****** Python file project directory :  {file_name} **********>>")
	
	# Open and read the content of the .py file
	with open(py_file, 'r') as file:
		content = file.read()
		# Append the content to the document
		doc.add_paragraph(content)
		doc.add_paragraph('\n')  # Add a newline for separation between files

# Define the source directory for the second block
source_directory_2 = r'C:\Users\OMCENTER\projects\mytools\droidlogin\droid\authenticate'

# Find all .py files in the second source directory
py_files_2 = glob.glob(os.path.join(source_directory_2, '*.py'))

# Loop through each .py file in the second source directory
for py_file in py_files_2:
	# Add the file name before the content
	file_name = os.path.basename(py_file)
	doc.add_paragraph(f"<<****** Python file app directory :  {file_name} **********>>")
	
	# Open and read the content of the .py file
	with open(py_file, 'r') as file:
		content = file.read()
		# Append the content to the document
		doc.add_paragraph(content)
		doc.add_paragraph('\n')  # Add a newline for separation between files

# Save the document
doc.save(destination_file)